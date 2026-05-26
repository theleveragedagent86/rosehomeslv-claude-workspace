"""Build Brian Esposito's Instagram engagement strategy & plan as a Word doc."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/ryanrose/Downloads/Claude/AI Clients/Brian Esposito/Brian_Esposito_IG_Engagement_Plan.docx"

NAVY = RGBColor(0x0B, 0x2A, 0x4A)
GRAY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p


def callout(label, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}  ")
    run.bold = True
    run.font.color.rgb = NAVY
    p.add_run(text)
    p.paragraph_format.space_after = Pt(8)


# ---- Cover ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Instagram Engagement Strategy & Build Plan")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = NAVY
title.paragraph_format.space_before = Pt(60)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Prepared for Brian Esposito  |  Branch Manager, Nations Lending")
run.font.size = Pt(13)
run.font.color.rgb = GRAY
sub.paragraph_format.space_after = Pt(4)

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date.add_run("April 2026")
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('"Show up where the talent already is. Add real value. Let the relationships do the recruiting."')
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = NAVY
tagline.paragraph_format.space_before = Pt(40)

doc.add_page_break()

# ---- Section 1 ----
h1("1. The Big Idea")
body(
    "Brian, you've spent 28+ years in Las Vegas building relationships, and the through-line of every "
    "chapter, PGA pro at TPC, branch leader at HFG, founder of Par for the Cure, and now Branch Manager "
    "at Nations Lending, is the same: you turn relationships into long-term friendships. That's already "
    "your competitive advantage. This plan is about pointing that strength at one specific goal: growing "
    "your Nations Lending branch by attracting loan officers from competing companies."
)
body(
    "We're calling the strategy value-provided marketing. Instead of cold-DMing LOs or running ads, you'll "
    "show up consistently in the comment sections of competitor loan officers across Vegas, Henderson, and "
    "Clark County, leaving tight 1-2 sentence comments (maximum 2) that add genuine insight. Not pitches. Not "
    '"check me out." Real value. Done daily, this builds name recognition with the exact people you want '
    "on your team. They start seeing 'Espo' show up everywhere with smart takes, and eventually one of "
    "them DMs you about a coffee."
)
callout("Why this works for you specifically:",
        "Right now your public content has zero recruiting angle, and that's whitespace. Most branch managers "
        "in Vegas are invisible to the LOs they want to hire. You can own that lane without changing your "
        "brand voice or compromising your relationship-first values.")

# ---- Section 2 ----
h1("2. Who We Built This Around")
body("Before we wrote a single line of this plan, we built a profile of you to make sure the system would sound like you, not like a robot. Here's what we anchored on:")
bullet("Brian J. Esposito  |  Branch Manager, Nations Lending  |  NMLS #1535781")
bullet("28+ years in Las Vegas. Former PGA TOUR / TPC Head Golf Pro at TPC Canyons (now TPC Las Vegas) and TPC Deere Run, home of the John Deere Classic.")
bullet('Founder of Par for the Cure (501c3), raised over $1.6 million over 20 years for breast cancer research in honor of your mother JoAnn. Annual October tournament at Bear\'s Best.')
bullet("Brand on Instagram: @espohomeloans (business account only). Nickname: Espo.")
bullet('Voice: warm, casual, story-driven, relationship-first. Aligns with the Nations Lending tagline "Home loans made human :)".')
bullet("Areas of expertise we'll lean on in your comments: mortgage education, Vegas market knowledge, branch leadership, golf credibility, and giving-back / community.")
body(
    "If anything in the profile above needs correcting, flag it before we go live, especially "
    "current branch address and team size."
)

# ---- Section 3 ----
h1("3. What You're Getting")
body(
    "We're building you two custom tools that run inside Claude Code (the same AI environment we'd help "
    "you set up). They're called skills — think of them like saved playbooks the AI runs on command. You "
    "type a short instruction, the system does the work."
)

h2("Tool 1: lender-comments  (your daily engagement engine)")
body(
    "This is the workhorse. You type 'run lender-comments' once a day, and the system opens a real "
    "Instagram browser session, rotates through 3 to 5 of your target competitor LOs, and engages with "
    "10 to 25 of their recent posts. For each post it likes the post, likes 5 to 10 of the top comments, "
    "and writes a 1-2 sentence value-add comment in your voice (hard cap at 2 sentences). It paces itself naturally to stay "
    "under Instagram's radar and stops immediately if it sees any rate-limit warning."
)
body("At the end of every session you get a clean log: which posts you engaged with, what you commented, what got skipped and why.")

h2("Tool 2: lender-research  (your target list builder)")
body(
    "Run this once a month or whenever your target list feels stale. It scours NMLS Consumer Access, "
    "Google, Instagram, and competitor company sites to find new Vegas-area mortgage LOs at companies "
    "other than Nations Lending. It verifies each candidate is active, posts regularly, and is genuinely "
    "in your market, then adds them to your target list with notes on their content style. You'll start "
    "with 15 to 25 verified accounts and grow from there."
)
body(
    "We're also integrating your MMI production data. You'll export a list of LOs in the $2M to $10M "
    "production sweet spot (PDF or Excel), and the system uses that to prioritize who to research and "
    "find on Instagram. That way you're not just engaging random LOs, you're engaging the ones who are "
    "producing enough to be worth recruiting but not so much that they've already got everything figured "
    "out. That's the sweet spot for growth."
)

# ---- Section 4 ----
h1("4. What a Typical Day Looks Like")
bullet("Morning (or whenever you have 10 minutes): open Claude Code, type 'run lender-comments'.")
bullet("The system opens Instagram in a browser window you can watch in real time.")
bullet("It rotates through 3 to 5 of your target accounts, opens their recent posts, and engages thoughtfully on 10 to 25 of them.")
bullet("Total runtime: roughly 15 to 30 minutes depending on how much it skips.")
bullet("You get a summary at the end: 'Engaged with 18 posts. Skipped 4 (2 political, 2 rate quotes). Left comments on all 18.'")
bullet("You can review every comment in the daily log and tweak the voice guidelines anytime if something feels off.")
body("If you want it fully hands-off, we can also set it up to run automatically every morning via Cowork's scheduler. You'll get a daily report without lifting a finger.")

# ---- Section 5 ----
h1("5. What Your Comments Will Sound Like")
body(
    "We dialed your voice in at warm professional peer crossed with mentor. Substantive, never salesy. "
    "Maximum 2 sentences, 1-2 only, hard cap. If it can be said in one sentence, say it in one. Lead with insight or a real question, not generic praise. Emoji allowed but "
    "sparingly, matching your :) brand energy."
)

h2("Three example comments")
callout("On a competitor LO's post about a tough VA loan they closed:",
        '"Love seeing VA closings done right. The hardest part is usually breaking the myth of VA lending."')
callout("On a competitor LO's post about coaching their new junior LO:",
        '"Coaching a new LO through their first year is the whole job. Consistency on the daily activities is what sticks the landing."')
callout("On a competitor LO's post about a Henderson neighborhood market update:",
        '"Henderson inventory has been wild this quarter. Cadence Hills and Inspirada are moving completely differently and most folks miss the nuance."')

h2("What you'll never see in your comments")
bullet('"DM me," "join my team," "check me out," "follow me"')
bullet("Any mention of Nations Lending")
bullet("Specific interest rates or APRs (compliance risk)")
bullet("Anything that bashes the competitor lender or their work")
bullet("Political, religious, or controversial social topics (system hard-skips these posts entirely)")
bullet("Generic praise or flattery without substance")

# ---- Section 6 ----
h1("6. Who We'll Engage With")
body(
    "Your target list is competitor mortgage LOs in Las Vegas, Henderson, and Clark County. We'll build "
    "out 15 to 25 verified accounts to start, organized by category so we can rotate intelligently:"
)
bullet("Independent mortgage brokers")
bullet("Retail branch loan officers (Rocket, UWM, loanDepot, CrossCountry, Guild, Fairway, Movement, Cardinal, Cornerstone, Supreme, PrimeLending, etc.)")
bullet("Other branch managers and team leaders (peer-to-peer recruiting whitespace)")
bullet("Wholesale account executives who interact with retail LOs in Vegas")
bullet("Emerging voices, newer LOs building their following")
body("Excluded: every Nations Lending account (your colleagues), and anyone you flag personally.")

# ---- Section 7 ----
h1("7. Setup at a Glance")
body(
    "When you give us the green light, we'll send you a separate setup walkthrough doc with screenshots. "
    "Here's the high-level so you know what's coming:"
)
bullet("Install Claude Code on your Mac or PC (free, ~5 minutes)")
bullet("Create or log into Cowork (the cloud environment that runs the browser automation)")
bullet("Install the Playwright MCP browser tool (one command, we'll walk you through it)")
bullet("Drop the two skill folders into your Claude directory")
bullet("Log into Instagram once in the Playwright browser; the session stays logged in after that")
bullet("Test run on a single account to confirm everything works")
body("Total setup time: roughly 30 to 45 minutes with us on a screen-share. You don't need any technical background.")

# ---- Section 8 ----
h1("8. What Success Looks Like")
body(
    "This is a slow-cooker, not a microwave. The honest expectation curve looks like this:"
)
bullet("Weeks 1 to 4: name recognition phase. Target LOs start seeing 'Espo' in their notifications regularly. No inbound yet, and that's normal.")
bullet("Weeks 4 to 12: warm familiarity. A few of them start liking your comments back, maybe DMing about something tactical. First 1 or 2 recruiting conversations.")
bullet("Months 3 to 6: inbound traction. The system has built enough surface area that interested LOs start reaching out about your branch organically.")
h2("Quick win: add a recruiting link to your Instagram bio")
body(
    "Right now your @espohomeloans profile doesn't have a link for LOs interested in joining your branch. "
    "Instagram lets you add multiple bio links. We recommend adding one that goes directly to your "
    "recruiting or 'join the team' page. That way, when an LO clicks through to your profile after seeing "
    "your comments, the path to start a conversation is already there. We'll help you set this up during "
    "our install session."
)

h2("The metrics that actually matter (not vanity metrics)")
bullet("Inbound DMs from competitor LOs")
bullet("Profile visits to @espohomeloans from accounts you commented on")
bullet("Recruiting conversations started")
bullet("LOs hired from this funnel (the real number)")

# ---- Section 9 ----
h1("9. Guardrails We Built In")
body("Because this runs on a real Instagram account using your real voice, we put hard rules in place:")
bullet("The system hard-skips political, religious, LGBTQ+, race, and police controversy posts. No exceptions.")
bullet("It will never quote a rate or make a pricing claim. Too much compliance risk.")
bullet("It will never mention Nations Lending in a comment.")
bullet("It paces itself between actions to stay under Instagram's spam thresholds.")
bullet('It stops immediately if Instagram shows any "Try Again Later" warning.')
bullet("Every comment is logged so you can review and course-correct.")
bullet('Default rule for the system: "When in doubt, skip the post."')

# ---- Section 10 ----
h1("10. Next Steps")
body("Here's what's happening next:")
bullet("You send over your MMI production list (PDF or Excel), filtered to the $2M to $10M sweet spot. We'll use that to build your starter target list of 15 to 25 verified LO Instagram accounts.")
bullet("Add a recruiting or 'join the team' link to your @espohomeloans Instagram bio (we'll help with this during setup).")
bullet("We build out both skills (lender-comments and lender-research), your comment voice guidelines, and the full target account list.")
bullet("We meet in person Thursday afternoon for a hands-on setup session: install Claude Code, Cowork, the browser automation tool, and get you logged in.")
bullet("We run a live test on a single target account so you can see exactly how it sounds before it goes live.")
body(
    "If you think of any specific competitor LOs you want prioritized, or anyone you want excluded, "
    "send those over along with the MMI list. Otherwise, we're good to start building."
)

# ---- Footer / contact ----
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = foot.add_run("Prepared by Ryan Rose  |  Rose Homes LV")
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = GRAY

doc.save(OUT)
print(f"Saved: {OUT}")
