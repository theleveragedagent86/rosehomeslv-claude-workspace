"""Build the First Responder Community Engagement strategy as a Word doc."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/ryanrose/Downloads/Claude/AI Clients/Brian Esposito/Brian_First_Responder_Engagement_Plan.docx"

NAVY = RGBColor(0x0B, 0x2A, 0x4A)
GRAY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)


def callout(label, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}  ")
    run.bold = True
    run.font.color.rgb = NAVY
    p.add_run(text)
    p.paragraph_format.space_after = Pt(8)


# ---- Cover ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("First Responder Community Engagement Plan")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = NAVY
title.paragraph_format.space_before = Pt(60)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Prepared for Brian Esposito  |  Branch Manager, Nations Lending")
run.font.size = Pt(13)
run.font.color.rgb = GRAY
sub.paragraph_format.space_after = Pt(4)

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date.add_run("April 2026")
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('"Show up as a neighbor. Acknowledge the work. Never reach for the business."')
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = NAVY
tagline.paragraph_format.space_before = Pt(40)

doc.add_page_break()

# ---- Section 1 ----
h1("1. The Big Idea")
body(
    "Brian, the lender plugin is your recruiting engine. This second plugin is something different and complementary. "
    "It is your community presence engine. The goal is to show up consistently in the comment sections of Las Vegas "
    "first responders, military and veterans, police, fire and EMS, and healthcare workers, with warm, genuine, "
    "1-2 sentence comments that acknowledge the work. Not pitches, not VA loan offers, not hero programs, not "
    "anything that sounds like a service provider. Just a recognizable Vegas neighbor showing up consistently."
)
body(
    "The downstream business benefit is real but indirect. When a Vegas firefighter eventually buys their first "
    "home, they will remember the guy who has been quietly showing up in their community comment threads for two "
    "years. When their cousin asks who to use for a VA loan, your name will come up. But the moment we reach for "
    "that benefit in a comment, the entire strategy collapses. The strategy works precisely because it is not "
    "transactional."
)
callout("Why this works for you specifically:",
        "You have 28+ years in this valley, Par for the Cure with $1.6M+ raised, and a former PGA Pro background "
        "that shows you know what showing up for community looks like. You are not a stranger reaching into a "
        "first responder community for leads. You are a recognizable Vegas community member who happens to be in "
        "mortgage. Those are very different identities, and this skill leans on the first one.")

# ---- Section 2 ----
h1("2. How This Differs From The Lender Plugin")
body(
    "The lender plugin and this plugin look similar mechanically (both run on @espohomeloans, both use the same "
    "browser automation, both engage on Instagram daily) but the voice and rules are very different. Here is the "
    "side-by-side:"
)

bullet("Lender plugin voice: industry peer crossed with mentor. Mortgage expertise on display. Implicit recruiting goal.")
bullet("Responder plugin voice: warm 28-year Vegas neighbor. Gratitude and acknowledgment. No business angle, ever.")
bullet("Lender comments: lead with insight or a real question.")
bullet("Responder comments: lead with acknowledgment of something specific in the post.")
bullet("Lender plugin: 1-2 reply comments per post when there is something additive.")
bullet("Responder plugin: default is no replies. Only reply when the tone is unambiguously safe and you have something genuine to add.")
bullet("Lender plugin: like 5-10 top comments per post.")
bullet("Responder plugin: like only 3-5 aligned top comments. First responder posts have wider tonal range and not all top comments are safe to like.")
bullet("Lender plugin: skip rate is moderate.")
bullet("Responder plugin: skip rate is high. When in doubt, skip. The downside risk of a misread comment is significant here.")

# ---- Section 3 ----
h1("3. Who We Will Engage With")
body(
    "Vegas, Henderson, and Clark County only. Out-of-state first responders are not in scope. We are targeting four "
    "categories with rough balance, 5-8 active accounts per category for a total of 20-30:"
)

h2("Military / Veterans")
bullet("Active duty service members (Nellis AFB, Creech AFB, Nevada Air and Army National Guard)")
bullet("Local veteran community accounts and VFW post members")

h2("Police / Law Enforcement")
bullet("LVMPD, Henderson PD, North LV PD, Boulder City PD, NV Highway Patrol, CCSD PD, UNLV PD")
bullet("Individual officers who post personally, not just department PR accounts")

h2("Fire / EMS")
bullet("Clark County Fire, Henderson Fire, LV Fire & Rescue, North LV Fire, Boulder City Fire")
bullet("Individual firefighters, paramedics, EMTs, K9 handlers")

h2("Healthcare")
bullet("Sunrise, UMC, Valley, Henderson Hospital, MountainView, Spring Valley, St Rose, Summerlin")
bullet("Nurses, doctors, frontline medical workers who post mix of work and community content")

body("How we find them: there is no MMI equivalent here, so the /responder-research skill mines department accounts, hashtag searches (#lvmpd, #vegasfire, #nellisafb, #vegasnurse, etc.), and bio keyword searches to discover individual accounts. Each candidate gets verified for Vegas geography, public account, regular activity, apolitical content mix, and comments enabled before being added to the target list.")

# ---- Section 4 ----
h1("4. What Your Comments Will Sound Like")
body(
    "Voice: warm 28-year Vegas neighbor. Genuine appreciation. Acknowledgment of specifics. NEVER service provider. "
    "Maximum 2 sentences, hard cap. One sentence is often better. If you cannot say something specific and "
    "genuine, skip the post."
)

h2("Six example comments by post type")

callout("On a Toys for Tots fire department drive:",
        '"This is the kind of thing that makes the Vegas community what it is. Thank you for showing up."')

callout("On a police K9 retirement post:",
        '"He earned every minute of this. Twelve years of work in those eyes."')

callout("On a firefighter posting about a training drill:",
        '"The daily reps that nobody sees are what makes the big calls go right. Respect."')

callout("On a 25-year retirement announcement:",
        '"Twenty-five years of this is no small thing. Enjoy the next chapter."')

callout("On a department fundraiser:",
        '"Pulling off events like this is its own full-time job. Hat tip to whoever ran point on it."')

callout("On a Memorial Day post (genuine remembrance, not partisan):",
        '"Remembering the names today. Thank you for keeping the watch."')

h2("What you will NEVER see in these comments")
bullet("Any mention of mortgages, lending, Nations Lending, VA loans, hero programs, or business")
bullet('Any service-provider phrases: "happy to help," "let me know if I can assist," "always here for you"')
bullet('"Stay safe out there, brother" or "thank you for your service" alone (overused, performative from a stranger)')
bullet("Any CTA, link, or @mention of your accounts")
bullet("Anything that sounds like prospecting")

# ---- Section 5 ----
h1("5. Guardrails (Stricter Than the Lender Plugin)")
body(
    "First responder content trips guardrails more often than mortgage industry content. The skill is built to skip "
    "aggressively rather than risk a misread comment. These posts are hard skips:"
)
bullet("Political content of any direction (police reform, military policy, immigration debates, election content)")
bullet("Active line-of-duty death posts (unless you personally knew them)")
bullet("Active critical incident posts (active shooter, mass casualty) until fully resolved")
bullet("Polarizing religious content")
bullet("Vaccine, mask, COVID controversy")
bullet("LGBTQ+ topics")
bullet("Race relations framed as debates")
bullet("Venting or complaints about the public, the department, the job")
bullet("Posts under 2 hours old with no engagement")
bullet("Anything where commenting could look opportunistic from a mortgage business account")

callout("Default rule:",
        "When in doubt, skip. There are plenty of safe posts to engage with. The downside of one wrong comment "
        "from your business account on a first responder post is significant.")

# ---- Section 6 ----
h1("6. Daily Cadence")
body(
    "This skill runs on the same @espohomeloans account as the lender plugin. To stay under Instagram's natural-"
    "engagement thresholds, combined daily volume should stay around 30 posts max."
)
bullet("Schedule the two skills at different times: e.g., lender-comments at 9 AM, responder-comments at 2 PM.")
bullet("If lender-comments has already run heavy (15+ posts) today, run responder-comments lighter (5-10 posts).")
bullet("If you only run one per day, alternate between the two.")
bullet("3-5 first responder accounts per session, mixing at least 2 different categories (don't engage all police on the same day).")
bullet("10-25 posts engaged per session as the upper target.")
bullet("0-1 replies per post (default: zero).")

# ---- Section 7 ----
h1("7. What Success Looks Like")
body(
    "This is even more of a slow burn than the lender plugin. The honest expectation curve:"
)
bullet("Months 1 to 3: pure presence. First responder targets see your name show up consistently in their notifications. No inbound, no measurable activity. This is the foundation.")
bullet("Months 3 to 6: warm familiarity. A few accounts start liking your comments back. You may get some follower growth from this audience. Still no direct business activity.")
bullet("Months 6 to 12: indirect referrals start. Someone's spouse asks who to use for a VA loan, your name comes up because their husband recognizes you from Instagram. A nurse colleague asks who closed their friend's first home cleanly.")
bullet("Year 2 and beyond: this audience becomes a quiet but reliable referral pipeline that compounds because it is built on real recognition, not interruption marketing.")

body("The metrics that actually matter (not vanity metrics):")
bullet("Inbound DMs from first responders or their family members")
bullet("Profile visits to @espohomeloans from first responder follower accounts")
bullet("Referral conversations that mention 'I see you on Instagram'")
bullet("VA loans / hero program closings sourced from this audience over time")

# ---- Section 8 ----
h1("8. What We Need From You")
body("Most of this is already in motion. Here is what is left:")
bullet("Confirm you want to run this in addition to the lender plugin (combined volume on the same account)")
bullet("Decide whether you want to schedule both skills automatically or run them manually each day")
bullet("Flag anyone you do NOT want targeted (specific officers, departments, individuals you have personal connections with that would be awkward)")
bullet("Approve the voice once we run a dry test (drafts only, no posting) on the first session")
body(
    "Once you give the green light, we install the plugin alongside the lender plugin on your @espohomeloans setup, "
    "run /responder-research to populate the initial 20-30 target list, and dry-run the first session for your "
    "review before anything goes live."
)

# ---- Footer ----
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = foot.add_run("Prepared by Ryan Rose  |  Rose Homes LV  |  ryan@rosehomeslv.com")
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = GRAY

doc.save(OUT)
print(f"Saved: {OUT}")
