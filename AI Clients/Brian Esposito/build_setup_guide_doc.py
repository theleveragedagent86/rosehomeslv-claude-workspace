"""Build Brian's step-by-step setup guide as a Word doc."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/ryanrose/Downloads/Claude/AI Clients/Brian Esposito/Brian_Setup_Guide.docx"
NAVY = RGBColor(0x0B, 0x2A, 0x4A)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1E, 0x7E, 0x34)

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def step(num, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{num}. ")
    r.bold = True
    p.add_run(text)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)


def bullet(text, indent=0.5):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)


def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(6)


def verify(text):
    p = doc.add_paragraph()
    r = p.add_run("✓ How to verify it worked:  ")
    r.bold = True
    r.font.color.rgb = GREEN
    p.add_run(text)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(10)


# ---- Cover ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Setup Guide")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = NAVY
title.paragraph_format.space_before = Pt(50)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Instagram Engagement Plugin  |  Brian Esposito")
r.font.size = Pt(13)
r.font.color.rgb = GRAY
sub.paragraph_format.space_after = Pt(30)

summary = doc.add_paragraph()
summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = summary.add_run("Total time: 30 to 45 minutes")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = NAVY

needs = doc.add_paragraph()
needs.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = needs.add_run("What you need: a computer, your @espohomeloans Instagram login, and the lender-engagement.zip file")
r.font.size = Pt(10)
r.font.color.rgb = GRAY
needs.paragraph_format.space_after = Pt(20)

doc.add_page_break()

# ---- Part 1 ----
h1("Part 1  |  Install Claude Code  (5 minutes)")
body("Claude Code is the AI tool that runs the plugin.")
step(1, "Go to https://claude.com/download in your browser")
step(2, 'Click "Download for Mac" (or Windows if you\'re on a PC)')
step(3, "Open the downloaded file and follow the installer")
step(4, "Once installed, open the Claude app")
step(5, "Sign in with your Claude account, or create one at https://claude.com if you don't have one yet. You'll need Claude Pro ($20/month) to start.")
verify("Open Claude and you see a chat window. That's it.")

# ---- Part 2 ----
h1("Part 2  |  Set Up Cowork  (10 minutes)")
body("Cowork is the cloud environment that lets Claude control a browser, so it can post on Instagram for you.")
step(1, "Go to https://cowork.ai in your browser")
step(2, 'Click "Sign Up" or "Get Started"')
step(3, "Sign in with the same Claude account from Part 1")
step(4, "Pick a plan, the entry-level plan is fine to start with")
step(5, "You'll land on the Cowork dashboard")
verify("You see a Cowork workspace with a sidebar and a chat area.")

# ---- Part 3 ----
h1("Part 3  |  Install the Plugin  (5 minutes)")
body("This loads both skills (/lender-comments and /lender-research) plus the browser automation into your Cowork workspace.")

h2("Option A: Install via zip upload (easiest)")
step(1, 'In Cowork, look for "Plugins" or "Extensions" in the settings or sidebar menu')
step(2, 'Click "Install Plugin" or "Upload Plugin"')
step(3, "Select the lender-engagement.zip file I sent you")
step(4, "Click Install and wait for confirmation")
step(5, "You should see lender-engagement listed as an installed plugin")

h2("Option B: Manual install (if Option A isn't available)")
step(1, "Unzip lender-engagement.zip on your computer")
step(2, "In Cowork, open the file explorer or terminal")
step(3, "Drag the unzipped lender-engagement folder into Cowork's ~/.claude/plugins/ directory")
step(4, "Restart Cowork")

verify("In the Cowork chat, type / and you see /lender-comments and /lender-research in the dropdown menu.")

# ---- Part 4 ----
h1("Part 4  |  Log Into Instagram  (5 minutes)")
body("The plugin uses a real browser to post on Instagram. You need to log in once so it stays logged in going forward.")
step(1, "In Cowork, type this exact message and send it:")
code("Use playwright-instagram to open a browser to https://www.instagram.com/accounts/login/")
step(2, "A Chrome window will open with the Instagram login page")
step(3, "Log in manually with your @espohomeloans credentials")
step(4, "If you use 2FA, complete the 2FA prompt")
step(5, "Once you're on the Instagram home feed, come back to Cowork and type:")
code("I'm logged in, you can continue")
verify("Your login is now saved. You won't need to log in again unless Instagram forces you to.")

# ---- Part 5 ----
h1("Part 5  |  Add Your Recruiting Link to Your Bio  (2 minutes)")
body("This is the link LOs will click after they see your comments and want to learn about your branch.")
step(1, "Open Instagram on your phone or computer")
step(2, "Go to your @espohomeloans profile")
step(3, 'Tap "Edit Profile"')
step(4, 'Scroll to "Links" (or "Website")')
step(5, 'Tap "Add External Link"')
step(6, 'Paste your recruiting page URL (the "join the team" page you showed me on your phone)')
step(7, 'Label it "Join the Team" or "Careers"')
step(8, "Save")
verify('Your @espohomeloans profile now shows a "Join the Team" link.')

# ---- Part 6 ----
h1("Part 6  |  Do a Test Run  (10 minutes)")
body("Let's run the skill on a SINGLE post first so you can see exactly what it does before going full volume.")
step(1, "In Cowork, type:")
code("/lender-comments pick one account and show me the comments you would write without posting them yet")
step(2, "Claude will open Instagram, pick one of your 14 target accounts, read recent posts, draft 1-2 sentence comments (maximum 2 sentences), and show them to you WITHOUT posting anything.")
step(3, "Review the drafts. Do they sound like you? Anything feel off?")
step(4, 'If anything needs adjusting, tell Claude exactly what to fix (e.g., "less formal, more like a peer")')
step(5, "When you're happy with the drafts, type:")
code("Great, now actually run the full routine")
step(6, "Claude will engage on 10 to 25 posts across 3 to 5 accounts, pacing itself naturally (about 15 to 30 minutes total)")
step(7, "At the end, you'll get a summary of what was posted")
verify("Check a couple of target accounts on Instagram. You should see comments from @espohomeloans on their recent posts.")

# ---- Part 7 ----
h1("Part 7  |  Set Up the Daily Schedule  (5 minutes, optional)")
body("If you want it to run every morning automatically:")
step(1, 'In Cowork, look for "Scheduler" or "Scheduled Tasks" in the settings')
step(2, 'Click "Create Scheduled Task" or "New Schedule"')
step(3, "Set frequency: Daily")
step(4, "Set time: whatever works for you. I recommend 9:00 AM PT on weekdays.")
step(5, "For the prompt, copy and paste exactly what's in the plugin's daily-schedule-prompt.md file.")
bullet("Find it in Cowork at: ~/.claude/plugins/lender-engagement/skills/lender-comments/references/daily-schedule-prompt.md")
step(6, "Save the schedule")
verify("Come back tomorrow morning and check the log at output/lender-comments/YYYY-MM-DD.md. If it's there, the schedule is working.")

# ---- Your daily workflow ----
h1("Your Daily Workflow After Setup")
h2("Option 1: Manual daily run (5 minutes)")
bullet("Open Cowork")
bullet('Type: "run lender-comments"')
bullet("Wait 15 to 30 minutes while it engages")
bullet("Review the log when done")

h2("Option 2: Fully automated")
bullet("The scheduled task runs every morning")
bullet("You get a daily log at output/lender-comments/YYYY-MM-DD.md")
bullet("Skim the log to make sure nothing looks off")

h2("Monthly: refresh the target list")
bullet("Send me your latest MMI production export (PDF or Excel)")
bullet("I'll run /lender-research with the MMI data to prioritize $2M-$10M producers")
bullet("Your target list gets updated with fresh, production-ranked LOs")

# ---- Troubleshooting ----
h1("What To Do If Something Breaks")

h2('"I don\'t see /lender-comments in the menu"')
body("The plugin didn't install. Uninstall and reinstall from Part 3.")

h2('"Instagram says I\'m logged out"')
body("Repeat Part 4 to log back in. Your session expired (this happens sometimes).")

h2('"Claude says Try Again Later"')
body("Instagram rate limited the account. Stop running it for 24 hours. Then resume with smaller sessions (5 to 10 posts instead of 10 to 25).")

h2('"The comments don\'t sound like me"')
body("Tell me exactly what's off. I'll update the voice guidelines in the plugin and send you a new version.")

h2('"A competitor LO moved to Nations Lending"')
body("Let me know their @ handle. I'll remove them from the target list.")

# ---- Footer ----
doc.add_paragraph()
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = contact.add_run("Call me if you get stuck. Don't push through it.")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = NAVY

foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Ryan Rose  |  ryan@rosehomeslv.com")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GRAY

doc.save(OUT)
print(f"Saved: {OUT}")
